import io

from docx import Document as _DocxDocument

from app.rag.parsers.base import ParsedDocument, ParserError


def parse_docx(content: bytes) -> ParsedDocument:
    try:
        doc = _DocxDocument(io.BytesIO(content))
    except Exception as e:
        raise ParserError(f"Failed to open DOCX: {e}") from e

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))

    text = "\n".join(parts)
    if not text:
        raise ParserError("DOCX appears empty")
    return ParsedDocument(text=text, metadata={})
